#!/usr/bin/env python3
"""
Simple affidavit generator - clones defendant tables as needed
"""

import json
import sys
import base64
from docx import Document
from copy import deepcopy
import io

def generate_affidavit(data):
    """Generate affidavit by cloning defendant tables"""
    
    # Load template
    template_path = 'Form_11_-_Affidavit_of_Service.docx'
    doc = Document(template_path)
    
    # Fill in header data
    # Table 0: Case number
    doc.tables[0].rows[1].cells[4].text = data['caseNumber']
    
    # Table 1: Claimant
    doc.tables[1].rows[0].cells[1].text = data['claimant']
    
    # Get the First Defendant table (Table 2) - we'll use this as template
    first_defendant_table = doc.tables[2]
    
    # Get number of defendants needed
    defendants = data['allDefendants']
    defendant_count = len(defendants)
    
    # Fill in First Defendant
    first_defendant_table.rows[0].cells[1].text = defendants[0]
    
    # If we have more than 1 defendant, we need to clone
    if defendant_count > 1:
        # Get the element and parent
        first_def_element = first_defendant_table._element
        parent = first_def_element.getparent()
        first_def_pos = list(parent).index(first_def_element)
        
        # Get the spacing paragraph after first defendant
        spacing_element = parent[first_def_pos + 1] if first_def_pos + 1 < len(parent) else None
        
        # Current position to insert new defendants
        insert_pos = first_def_pos + 2  # After First Defendant and its spacing
        
        # Clone for defendants 2, 3, 4, etc.
        ordinals = ['Second', 'Third', 'Fourth', 'Fifth', 'Sixth']
        
        for i in range(1, defendant_count):
            if i >= 6:  # Max 6 defendants
                break
                
            # Clone the first defendant table
            new_table_elem = deepcopy(first_def_element)
            
            # Clone spacing if it exists
            if spacing_element is not None:
                new_spacing_elem = deepcopy(spacing_element)
            
            # Insert the new table
            parent.insert(insert_pos, new_table_elem)
            insert_pos += 1
            
            # Insert spacing
            if spacing_element is not None:
                parent.insert(insert_pos, new_spacing_elem)
                insert_pos += 1
    
    # Now we need to reload to access the new tables
    # Save to bytes and reload
    temp_buffer = io.BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)
    doc = Document(temp_buffer)
    
    # Now fill in all defendant tables with correct labels and values
    ordinals = ['First', 'Second', 'Third', 'Fourth', 'Fifth', 'Sixth']
    
    for i in range(defendant_count):
        table_idx = 2 + i  # Defendant tables start at index 2
        if table_idx < len(doc.tables):
            # Update label
            doc.tables[table_idx].rows[0].cells[0].text = f"{ordinals[i]} Defendant"
            # Update value
            doc.tables[table_idx].rows[0].cells[1].text = defendants[i]
    
    # Delete any extra defendant tables (the original Third-Sixth if they still exist)
    # Start from the end to avoid index shifting
    tables_to_check = len(doc.tables)
    for idx in range(tables_to_check - 1, 1, -1):
        if idx < len(doc.tables):
            table = doc.tables[idx]
            if len(table.rows) > 0:
                cell_text = table.rows[0].cells[0].text
                # If it's an empty defendant table (Third, Fourth, Fifth, Sixth with bracket fields)
                if ('Third Defendant' in cell_text or 'Fourth Defendant' in cell_text or 
                    'Fifth Defendant' in cell_text or 'Sixth Defendant' in cell_text):
                    # Check if it still has bracket field (meaning it wasn't filled)
                    if '[Defendant' in table.rows[0].cells[1].text:
                        tbl = table._element
                        tbl.getparent().remove(tbl)
    
    # Fill in service statement with correct defendant info
    # Find which defendant this affidavit is for
    current_defendant = data['defendantName']
    defendant_index = defendants.index(current_defendant)
    defendant_ordinal = f"{ordinals[defendant_index]} Defendant"
    
    # Find and update service statement table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'duly serve' in cell.text:
                    cell.text = cell.text.replace('[Name]', current_defendant)
                    cell.text = cell.text.replace('[Defendant]', defendant_ordinal)
                    cell.text = cell.text.replace('[Date]', '')
                    cell.text = cell.text.replace('[time am/pm]', '')
                    cell.text = cell.text.replace('[Place]', '')
                    cell.text = cell.text.replace('[Name of process]', 'General Procedure Claim')
    
    # Save to buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    return output_buffer.getvalue()

# Read input from stdin
input_data = json.loads(sys.stdin.read())

try:
    # Generate affidavit
    docx_bytes = generate_affidavit(input_data)
    
    # Return base64 encoded
    result = {
        'success': True,
        'data': base64.b64encode(docx_bytes).decode('utf-8')
    }
    print(json.dumps(result))
    
except Exception as e:
    result = {
        'success': False,
        'error': str(e)
    }
    print(json.dumps(result))
    sys.exit(1)